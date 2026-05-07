#!/usr/bin/env python3
"""
MD to Word Converter with Template Support & Better Code Formatting
"""
from fastapi import FastAPI, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn
import io
import re
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class FileData(BaseModel):
    name: str
    content: str

class MDToWordRequest(BaseModel):
    title: str = "Documentación"
    files: list[FileData]
    template: str = "default"

# Enhanced templates with better code formatting
TEMPLATES = {
    "default": {
        "title_font": "Segoe UI",
        "title_size": 28,
        "title_color": RGBColor(99, 102, 241),
        "title_bold": True,
        "heading1_font": "Segoe UI",
        "heading1_size": 18,
        "heading1_color": RGBColor(51, 51, 51),
        "heading1_bold": True,
        "heading2_size": 14,
        "heading2_color": RGBColor(68, 68, 68),
        "body_font": "Segoe UI",
        "body_size": 11,
        "code_bg": "F3F4F6",
        "code_border": "D1D5DB",
        "code_font_size": 8.5,
        "code_line_spacing": 1.0,
        "table_header_bg": "6366F1",
        "table_header_color": "FFFFFF",
        "code_padding": 4,
        "code_margin": 2
    },
    "bc-digital": {
        "title_font": "Poppins",
        "title_size": 36,
        "title_color": RGBColor(0, 0, 0),
        "title_bold": True,
        "heading1_font": "Poppins",
        "heading1_size": 14,
        "heading1_color": RGBColor(0, 0, 0),
        "heading1_bold": True,
        "heading2_size": 12,
        "heading2_color": RGBColor(0, 0, 0),
        "body_font": "Poppins",
        "body_size": 11,
        "code_bg": "F8F9FA",
        "code_border": "E5E7EB",
        "code_font_size": 8.5,
        "code_line_spacing": 1.0,
        "table_header_bg": "6366F1",
        "table_header_color": "FFFFFF",
        "code_padding": 4,
        "code_margin": 2
    },
    "bc-digital-code": {
        "title_font": "Poppins",
        "title_size": 36,
        "title_color": RGBColor(0, 0, 0),
        "title_bold": True,
        "heading1_font": "Poppins",
        "heading1_size": 14,
        "heading1_color": RGBColor(0, 0, 0),
        "heading1_bold": True,
        "heading2_size": 12,
        "heading2_color": RGBColor(0, 0, 0),
        "body_font": "Poppins",
        "body_size": 11,
        "code_bg": "1E293B",
        "code_border": "334155",
        "code_color": "E2E8F0",
        "code_font_size": 9,
        "code_line_spacing": 1.1,
        "table_header_bg": "6366F1",
        "table_header_color": "FFFFFF",
        "code_padding": 6,
        "code_margin": 4
    },
    "modern-purple": {
        "title_font": "Segoe UI",
        "title_size": 32,
        "title_color": RGBColor(99, 102, 241),
        "title_bold": True,
        "heading1_font": "Segoe UI",
        "heading1_size": 20,
        "heading1_color": RGBColor(51, 51, 51),
        "heading1_bold": True,
        "heading2_size": 16,
        "heading2_color": RGBColor(68, 68, 68),
        "body_font": "Segoe UI",
        "body_size": 11,
        "code_bg": "1E1E2E",
        "code_border": "4B5563",
        "code_color": "E0E0E0",
        "code_font_size": 9,
        "code_line_spacing": 1.1,
        "table_header_bg": "6366F1",
        "table_header_color": "FFFFFF",
        "code_padding": 6,
        "code_margin": 4
    }
}

def get_template(name):
    return TEMPLATES.get(name, TEMPLATES["default"])

def clean_code(code):
    """Remove blank lines from code (but keep meaningful structure)"""
    lines = code.split('\n')
    cleaned = []
    prev_empty = False
    
    for line in lines:
        # Check if line is empty or only whitespace
        is_empty = len(line.strip()) == 0
        
        # Only skip if we have consecutive empty lines
        if is_empty and prev_empty:
            continue
        prev_empty = is_empty
        
        # Remove trailing whitespace but keep the line
        cleaned.append(line.rstrip())
    
    # Remove leading/trailing empty lines
    while cleaned and not cleaned[0]:
        cleaned.pop(0)
    while cleaned and not cleaned[-1]:
        cleaned.pop()
    
    return '\n'.join(cleaned)

def add_code_block(doc, code, language="", template=None):
    """Add a nicely formatted code block with better readability"""
    if template is None:
        template = get_template("default")
    
    # Clean the code - remove consecutive blank lines
    code = clean_code(code)
    
    # Calculate padding values
    code_padding = template.get("code_padding", 4)
    code_margin = template.get("code_margin", 2)
    
    # Create a table for better code box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    
    # Set cell background
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), template.get("code_bg", "F3F4F6"))
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    
    # Remove table borders
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    # Set cell margins
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', str(code_padding * 20)), ('bottom', str(code_padding * 20)), 
                                       ('left', str(code_margin * 20)), ('right', str(code_margin * 20))]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), margin_value)
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)
    
    # Add code content line by line
    code_lines = code.split('\n')
    for i, line in enumerate(code_lines):
        para = cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = template.get("code_line_spacing", 1.0)
        
        # Add language tag on first line if provided
        if i == 0 and language:
            lang_run = para.add_run(f"// {language.upper()} ")
            lang_run.font.name = 'Consolas'
            lang_run.font.size = Pt(template.get("code_font_size", 9) - 1)
            lang_run.font.color.rgb = RGBColor(100, 116, 139)
            lang_run.font.italic = True
        
        # Add code line
        code_run = para.add_run(line if line else " ")
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(template.get("code_font_size", 9))
        code_color = template.get("code_color", RGBColor(40, 40, 40))
        if isinstance(code_color, str):
            code_color = RGBColor(int(code_color[0:2], 16), int(code_color[2:4], 16), int(code_color[4:6], 16))
        code_run.font.color.rgb = code_color
    
    # Remove default empty paragraph
    if cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]._element
        p.getparent().remove(p)

def add_styled_heading(doc, text, level, template=None):
    """Add a heading with template styling"""
    if template is None:
        template = get_template("default")
    
    p = doc.add_heading(level=level)
    run = p.add_run(text)

    if level == 1:
        run.font.name = template.get("heading1_font", "Segoe UI")
        run.font.size = Pt(template.get("heading1_size", 18))
        run.font.color.rgb = template.get("heading1_color", RGBColor(51, 51, 51))
        run.font.bold = template.get("heading1_bold", True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.name = template.get("heading1_font", "Segoe UI")
        run.font.size = Pt(template.get("heading2_size", 14))
        run.font.color.rgb = template.get("heading2_color", RGBColor(68, 68, 68))
        run.font.bold = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)

    return p

def add_styled_table(doc, rows, template=None):
    """Add a nicely formatted table"""
    if template is None:
        template = get_template("default")
    
    if not rows:
        return

    cols = len([c.strip() for c in rows[0].split('|') if c.strip()])
    if cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    header_bg = template.get("table_header_bg", "6366F1")
    header_color = template.get("table_header_color", "FFFFFF")

    for row_idx, row_data in enumerate(rows):
        cells = [c.strip() for c in row_data.split('|') if c.strip()]
        for col_idx, cell_text in enumerate(cells[:cols]):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text

            if row_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(int(header_color[0:2], 16), int(header_color[2:4], 16), int(header_color[4:6], 16))
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), header_bg)
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
            else:
                if row_idx % 2 == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'F8F8F8')
                    shd.set(qn('w:val'), 'clear')
                    tcPr.append(shd)

def process_markdown(content):
    """Process markdown content"""
    lines = content.split('\n')
    elements = []
    i = 0

    while i < len(lines):
        line = lines[i]

        if line.startswith('#### '):
            elements.append(('heading4', line[5:]))
        elif line.startswith('### '):
            elements.append(('heading3', line[4:]))
        elif line.startswith('## '):
            elements.append(('heading2', line[3:]))
        elif line.startswith('# '):
            elements.append(('heading1', line[2:]))
        elif line.strip().startswith('```'):
            code_lines = []
            language = line.strip()[3:].strip()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append(('code', '\n'.join(code_lines), language))
        elif line.strip().startswith('|'):
            table_rows = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|') and '---' not in lines[i]:
                table_rows.append(lines[i])
                i += 1
            elements.append(('table', table_rows))
            continue
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            elements.append(('ul', line.strip()[2:]))
        elif re.match(r'^\d+\. ', line.strip()):
            match = re.match(r'^(\d+)\. (.+)', line.strip())
            if match:
                elements.append(('ol', match.group(2)))
        elif line.strip() in ['---', '***', '___']:
            elements.append(('hr',))
        elif line.strip() == '':
            elements.append(('empty',))
        elif line.strip():
            elements.append(('para', line))

        i += 1

    return elements

@app.post("/api/tools/md-to-word")
def md_to_word(request: MDToWordRequest):
    """Convert multiple MD files to Word document with template"""

    doc = Document()
    doc.core_properties.title = request.title

    template = get_template(request.template)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading(level=0)
    title_run = title.add_run(request.title)
    title_run.font.name = template.get("title_font", "Segoe UI")
    title_run.font.size = Pt(template.get("title_size", 28))
    title_run.font.color.rgb = template.get("title_color", RGBColor(99, 102, 241))
    title_run.font.bold = template.get("title_bold", True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from datetime import datetime
    date_run = subtitle.add_run(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(150, 150, 150)
    date_run.font.italic = True

    doc.add_paragraph()

    for file_info in request.files:
        filename = file_info.name
        content = file_info.content

        file_title = doc.add_heading(level=1)
        file_run = file_title.add_run(f"📄 {filename.replace('.md', '')}")
        file_run.font.size = Pt(18)
        file_run.font.color.rgb = RGBColor(60, 60, 60)
        file_title.paragraph_format.space_before = Pt(16)
        file_title.paragraph_format.space_after = Pt(10)

        elements = process_markdown(content)

        for elem in elements:
            if elem[0] == 'heading1':
                add_styled_heading(doc, elem[1], 1, template)
            elif elem[0] == 'heading2':
                add_styled_heading(doc, elem[1], 2, template)
            elif elem[0] == 'heading3':
                add_styled_heading(doc, elem[1], 3, template)
            elif elem[0] == 'heading4':
                add_styled_heading(doc, elem[1], 4, template)
            elif elem[0] == 'code':
                add_code_block(doc, elem[1], elem[2], template)
            elif elem[0] == 'table':
                add_styled_table(doc, elem[1], template)
            elif elem[0] == 'ul':
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(elem[1])
                run.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
            elif elem[0] == 'ol':
                p = doc.add_paragraph(style='List Number')
                run = p.add_run(elem[1])
                run.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
            elif elem[0] == 'hr':
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(10)
                run = p.add_run('─' * 70)
                run.font.color.rgb = RGBColor(200, 200, 200)
            elif elem[0] == 'empty':
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
            elif elem[0] == 'para':
                p = doc.add_paragraph()
                run = p.add_run(elem[1])
                run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(4)

        if len(request.files) > 1:
            doc.add_page_break()

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return Response(
        content=doc_buffer.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{request.title}.docx"'}
    )

@app.get("/health")
def health():
    return {"status": "ok", "service": "md-to-word"}

if __name__ == "__main__":
    print("Starting MD to Word server with cleaner code blocks on port 8001...")
    uvicorn.run(app, host="0.0.0.0", port=8001, log_level="error")
